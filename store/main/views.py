import os
import tempfile
import tkinter
import tkinter.filedialog
from os.path import exists

from django.core.files.base import ContentFile
from django.http import HttpResponse, Http404, HttpResponseRedirect
from django.shortcuts import render, redirect
from docx import Document as Act
from docxtpl import *
from .models import *
from .forms import *
from django.views.generic import UpdateView, DeleteView
from datetime import datetime
from django.core.files import File

DATA = {}
ITEMS = {}


# ОТОБРАЖЕНИЕ ГЛАВНОЙ СТРАНИЦЫ
def inform_func(request):
    data = {'title': 'Главная страница',}
    return render(request, 'main/inform.html', data)


# ОТОБРАЖЕНИЕ ТАБЛИЦЫ С ЕД. ИЗМЕРЕНИЯ
def unit_func(request):
    unit = Unit.objects.all()

    return render(request, 'main/unit.html', {'unit': unit})


# ДОБАВЛЕНИЕ ЕДИНИЦЫ ИЗМЕРЕНИЯ
def add_unit_func(request):
    error = ''
    title = 'Форма добавления'
    if request.method == 'POST':
        form = UnitForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('unit')
        else:
            error = 'Форма была неверно заполнена.'
    form = UnitForm()
    data = {
        'form': form,
        'error': error,
        'title': title}
    return render(request, 'main/unit_add_form.html', data)


# ОТОБРАЖЕНИЕ С ТАБЛИЦЕЙ МЕСТ ХРАНЕНИЯ
def place_func(request):
    place = Place.objects.all()

    return render(request, 'main/place.html', {'place': place})

# ДОБАВЛЕНИЕ НОВОГО МЕСТА ХРАНЕНИЯ
def add_place_func(request):
    error = ''
    title = 'Форма добавления'
    if request.method == 'POST':
        form = PlaceForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('place')
        else:
            error = 'Форма была неверно заполнена.'
    form = PlaceForm()
    data = {
        'form': form,
        'error': error,
        'title': title}
    return render(request, 'main/place_add_form.html', data)


# ОТОБРАЖЕНИЕ И ФИЛЬТРАЦИЯ СТРАНИЦЫ "ПОСТАВЩИКИ"
def providers_func(request):
    search_query = request.GET.get('search', '').upper()
    providers = Provider.objects.all()
    key_list = []
    message = ''
    if search_query:
        for i in range(len(providers)):
            name = providers[i].provider_name.upper()
            key = providers[i].id
            if name.find(search_query) != -1:
                key_list.append(key)
        if len(key_list) == 0:
            message = 'Поставщик не найден. Список доступных поставщиков:'

    return render(request, 'main/providers_inform.html', {'providers': providers, 'key': key_list, 'message': message})

# РЕДАКТИРОВАНИЕ ПОСТАВЩИКА
class ProvidersUpdatelView(UpdateView):
    model = Provider
    template_name = 'main/provider_add_form.html'
    form_class = ProviderForm

# УДАЛЕНИЕ ПОСТАВЩИКА
def delete_provider_func(request, id=None):
    error_delete = ''
    arr = []
    providers = Provider.objects.all()
    providet_to_delete = Provider.objects.get(id=id)
    try:
        providet_to_delete.delete()
        return render(request, 'main/providers_inform.html', {'providers': providers, 'error_delete': error_delete})
    except models.ProtectedError:
        products = Product.objects.all()
        for p in products:
            if p.provider == providet_to_delete:
                error_delete = 'Нельзя удалить поставщика, так как на него  ссылаются следующие товары: '
                arr.append(p.product_name)
        return render(request, 'main/providers_inform.html',
                      {'providers': providers, 'error_delete': error_delete, 'arr': arr})

# ДОБАВЛЕНИЕ НОВОГО ПОСТАЩИКА
def add_provider_func(request):
    error = ''
    title = 'Форма добавления нового поставщика'
    if request.method == 'POST':
        form = ProviderForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('providers')
        else:
            error = 'Форма была неверно заполнена.'
    form = ProviderForm()
    data = {
        'form': form,
        'error': error,
        'title': title,
    }
    return render(request, 'main/provider_add_form.html', data)


# ОТОБРАЖЕНИЕ И ФИЛЬТРАЦИЯ СТРАНИЦЫ "КЛИЕНТЫ"
def clients_func(request):
    search_query = request.GET.get('search', '').upper()
    clients = Client.objects.all()
    key_list = []
    message = ''

    if search_query:
        for i in range(len(clients)):
            name = clients[i].client_name.upper()
            key = clients[i].id
            if name.find(search_query) != -1:
                key_list.append(key)
        if len(key_list) == 0:
            message = 'Клиент не найден. Список доступных клиентов:'

    return render(request, 'main/clients_inform.html', {'clients': clients, 'key': key_list, 'message': message})

# РЕДАКТИРОВАНИЕ КЛИЕНТА
class ClientsUpdatelView(UpdateView):
    model = Client
    template_name = 'main/client_add_form.html'
    form_class = ClientForm

# УДАЛЕНИЕ КЛИЕНТА
def delete_client_func(request, id=None):
    clients = Client.objects.all()
    client_to_delete = Client.objects.get(id=id)
    try:
        client_to_delete.delete()
        return render(request, 'main/clients_inform.html', {'clients': clients})
    except models.ProtectedError:
        error_delete = 'Нельзя удалить клиента, так как на него ссылаются различные документы. Обратитесь к администратору. '
        return render(request, 'main/clients_inform.html', {'clients': clients, 'error_delete': error_delete})

# ДОБАВЛЕНИЕ КЛИЕНТА
def add_client_func(request):
    error = ''
    title = 'Форма добавления нового клиента'
    if request.method == 'POST':
        form = ClientForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('clients')
        else:
            error = 'Форма была неверно заполнена.'
    form = ClientForm()
    data = {
        'form': form,
        'error': error,
        'title': title}
    return render(request, 'main/client_add_form.html', data)

# ОТОБРАЖЕНИЕ СТРАНИЦЫ "ТОВАРЫ" И ФИЛЬТРАЦИЯ
def products_func(request):
    search_query = request.GET.get('search', '').upper()

    products = Product.objects.all()
    key_list = []
    message = ''

    if search_query:
        for i in range(len(products)):
            name = products[i].product_name.upper()
            key = products[i].id
            if name.find(search_query) != -1:
                key_list.append(key)
        if len(key_list) == 0:
            message = 'Товар не найден. Список доступных товаров:'

    return render(request, 'main/products_inform.html', {'products': products, 'key': key_list, 'message': message})

# РЕДАКТИРОВАНИЕ ТОВАРА
class ProductsUpdatelView(UpdateView):
    model = Product
    template_name = 'main/product_add_form.html'
    form_class = ProductForm

# УДАЛЕНИЕ ТОВАРА
def delete_product_func(request, id=None):
    products = Product.objects.all()
    product_to_delete = Product.objects.get(id=id)
    try:
        product_to_delete.delete()
        return render(request, 'main/products_inform.html', {'products': products})
    except models.ProtectedError:
        error_delete = 'Нельзя удалить клиента, так как на него ссылаются различные документы. Обратитесь к администратору. '
        return render(request, 'main/products_inform.html', {'products': products, 'error_delete':error_delete})

# ДОБАВЛЕНИЕ ТОВАРА
def add_product_func(request):
    error = ''
    title = 'Форма добавления нового товара'

    if request.method == 'POST':
        form = ProductForm(request.POST)
        if form.is_valid():
            if form.fields['valid'] == None:
                form.fields['valid'] = ''
            form.save()
            return redirect('products')
        else:
            error = 'Форма была неверно заполнена.'

    form = ProductForm()
    data = {
        'form': form,
        'error': error,
        'title': title,
    }
    return render(request, 'main/product_add_form.html', data)


# ЗАГРУЗКА И ПРОСМОТР БЛАНКОВ ЗАКАЗА
def my_view(request):
    message = 'Загрузите бланк заказа'
    client = ''

    if request.method == 'POST':
        form = BlankForm(request.POST, request.FILES)
        if form.is_valid():
            client = form.cleaned_data.get('client')
            newdoc = Blank(status='Новый', docfile=request.FILES['docfile'], client=client,
                           date=str(datetime.today().strftime('%d.%m.%Y')))
            newdoc.save()
            return redirect('my-view')
        else:
            message = 'The form is not valid. Fix the following error:'
    else:
        form = BlankForm()

    documents = Blank.objects.all()

    context = {'documents': documents, 'form': form, 'message': message, 'client': client}
    return render(request, 'main/list.html', context)

# ИЗМЕНЕНИЕ СТАТУСА СБОРКИ ЗАКАЗА
def my_view_status(request, id=None):
    documents = Blank.objects.all()
    client = ''
    message = 'Загрузите бланк заказа'

    if (request.GET.get('next')):
        blank_to_change = Blank.objects.get(id=id)
        if blank_to_change.status == 'Новый':
            blank_to_change.status = 'Начат'
            blank_to_change.save()
        elif blank_to_change.status == 'Начат':
            blank_to_change.status = 'Закончен'
            blank_to_change.save()

    if request.method == 'POST':
        form = BlankForm(request.POST, request.FILES)
        if form.is_valid():
            client = form.cleaned_data.get('client')
            newdoc = Blank(status='Новый', docfile=request.FILES['docfile'], client=client,
                           date=str(datetime.today().strftime('%d.%m.%Y')))
            newdoc.save()
            return redirect('my-view')
        else:
            message = 'The form is not valid. Fix the following error:'
    else:
        form = BlankForm()

    context = {'documents': documents, 'form': form, 'message': message, 'client': client}

    return render(request, 'main/list.html', context)


def delete_blank_func(request, id=None):
    blank_to_delete = Blank.objects.get(id=id)
    blank_to_delete.delete()

    message = 'Загрузите бланк заказа'
    client = ''

    if request.method == 'POST':
        form = BlankForm(request.POST, request.FILES)
        if form.is_valid():
            client = form.cleaned_data.get('client')
            newdoc = Blank(status='Новый', docfile=request.FILES['docfile'], client=client,
                           date=str(datetime.today().strftime('%d.%m.%Y')))
            newdoc.save()
            return redirect('my-view')
        else:
            message = 'The form is not valid. Fix the following error:'
    else:
        form = BlankForm()

    documents = Blank.objects.all()

    context = {'documents': documents, 'form': form, 'message': message, 'client': client}

    return render(request, 'main/list.html', context)

# СКАЧИВАНИЕ ДОКУМЕНТОВ НА КОМПЬЮТЕР
def download_docx(request, path):
    print(path)
    file_path = os.path.join(settings.MEDIA_ROOT, path)
    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/vnd.ms-excel")
            response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
            # context={'response': response}
            return response
    raise Http404


C = 0
IDNUMACT = 0
IDNUMOPIS = 0
CO = 0
# ОТКРЫТИЕ ВКЛАДКИ "СФОРМИРОВАТЬ ДОКУМЕНТЫ"
def docShow_func(request):
    error = ''
    cols = ''
    cols_opis = ''
    numact = ''
    numopis = ''


    form = NumActForm(request.POST)
    if form.is_valid():
        global C, IDNUMACT
        C = 0
        i = 0

        numact = form.cleaned_data.get('numact')
        numact_object = NumAct(numact=numact)
        numact_object.save()

        IDNUMACT = numact_object.id
        cols = form.cleaned_data.get('cols')

        while i < int(cols):
                i = i + 1
                C = i


    form_opis = NumOpisForm(request.POST)
    if form_opis.is_valid():
        global CO, IDNUMOPIS
        CO = 0
        i = 0
        numopis = form_opis.cleaned_data.get('numopis')
        numopis_object = NumOpis(numopis=numopis)
        numopis_object.save()
        IDNUMOPIS = numopis_object.id
        cols_opis = form_opis.cleaned_data.get('cols_opis')
            # while i < int(cols_opis):
            #     i = i + 1
            #     CO = i
        CO = int(form_opis.cleaned_data.get('cols_opis'))


    data = {'title': 'Документы', 'form_opis': form_opis, 'numact': numact, 'cols': cols, 'form': form,
            'cols_opis': cols_opis, 'numopis': numopis, 'error': error}
    return render(request, 'main/docs.html', data)

# ЗАПОЛНЕНИЕ ФОРМЫ АКТА ПРИЕМА-ПЕРЕДАЧИ
def form_act(request):
    global C, IDNUMACT

    object = NumAct.objects.get(id=IDNUMACT)
    actnum = str(object.numact)
    actdate = ''
    company_ = ''
    post_ = ''
    name_ = ''
    contractnum = ''
    contractdate = ''
    docname = ''
    forms = []
    products_in_act = []

    error = ''
    error_prod = ''

    for i in range(C):
        forms.append(i)
        products_in_act.append(i)

    form = ActForm(request.POST or None)
    myform = ClientCompanyAddForm(request.POST or None)

    for i in range(len(forms)):
        prefix = 'form' + str(i)
        forms[i] = ProductsInActForm(request.POST or None, prefix=prefix)

        if form.is_valid() and myform.is_valid() and forms[i].is_valid():
            actdate = form.cleaned_data.get('act_date')
            company_ = myform.cleaned_data.get('client')
            post_ = form.cleaned_data.get('post')
            name_ = form.cleaned_data.get('name')
            contractnum = form.cleaned_data.get('contract_num')
            contractdate = form.cleaned_data.get('contract_date')
            docname = form.cleaned_data.get('doc_name')

            p = forms[i].cleaned_data.get('product')
            a = forms[i].cleaned_data.get('amount')

            p_for_check = Product.objects.get(product_name=p.product_name)
            print(p_for_check.amount)
            if int(p_for_check.amount) < int(a):
                error_prod = 'Не хватает товара.'

            products_in_act[i] = ProductsInAct(numact=NumAct.objects.get(id=IDNUMACT), product=p, amount=a)
            products_in_act[i].save()

        else:
            error = 'Форма была неверно заполнена.'

    context = {
        'form': form, 'myform': myform, 'forms': forms, 'C': C,
        'actnum': actnum,
        'actdate': actdate,
        'company_': company_, 'post_': post_, 'name_': name_,
        'contractnum': contractnum, 'contractdate': contractdate, 'docname': docname,
        'error': error, 'error_prod': error_prod,
        'products_in_act': products_in_act
    }

    print(context)

    if context['company_'] != '' and error_prod == '':
        make_act(context)

    return render(request, 'main/act_sale_add_form.html', context)

# СОЗДАНИЕ АКТА ПРИЕМА-ПЕРЕДАЧИ
def make_act(context):
    product_in_act = context['products_in_act'].copy()

    doc = Act('shablon.docx')
    str2 = '2. Принятый Покупателем товар обладает качеством и ассортиментом, соответствующим требованиям ' + \
           'Договора. Товар поставлен в установленные в Договоре сроки. Покупатель не имеет никаких претензий' + \
           ' к принятому товару. '
    str3 = '3. Настоящий Акт составлен в двух экземплярах, имеющих равную юридическую силу, по одному ' + \
           'экземпляру для каждой из Сторон, и является неотъемлемой частью Договора между Сторонами.'

    par = [p for p in doc.paragraphs]

    par[0].add_run(context['actnum'])
    if context['actdate'] == '':
        nowdate = str(datetime.today().strftime('"%d"   %m    %Yг.'))
    else:
        nowdate = context['actdate']
    par[1].add_run(nowdate)

    client_company = str(context['company_'])
    client_status = context['post_']
    client_name = context['name_']
    strPar = 'ООО "Технотекс", в лице генерального директора Плюшкина Андрея Степановича, действующего на ' + '' \
                                                                                                              'основании Устава, именуемый в дальнейшем Продавец, с одной стороны и ' + client_company + \
             ', в лице ' + client_status + ' ' + client_name + ', именуемый в дальнейшем Покупатель, ' + \
             'с другой стороны, составили Акт о нижеследующем:'

    doc.add_paragraph(strPar)

    dogovor_number = context['contractnum']
    dogovor_date = context['contractdate']
    doc.add_paragraph('1. В соответствии с условиями Договора, заключенного между Сторонами №' + dogovor_number +
                      ' от ' + dogovor_date + ' Поставщик поставляет, а Покупатель принимает Товар' +
                      ' следующего ассортимента и количества:')

    # Количество rows вводится
    r = len(product_in_act) + 2
    sum_of_order = float(0)

    table = doc.add_table(rows=r, cols=6)
    table.style = 'Table style'
    table.cell(0, 0).text = '№ '
    table.cell(0, 1).text = 'Наименование товара'
    table.cell(0, 2).text = 'Количество товара '
    table.cell(0, 3).text = 'Единица измерения'
    table.cell(0, 4).text = 'Цена за единицу товара '
    table.cell(0, 5).text = 'Сумма, включая НДС '
    table.cell(r - 1, 1).text = 'Итого: '

    if len(product_in_act) >= 1:
        for row in range(r):
            for col in range(6):
                if row != 0 and row != r - 1:
                    cell = table.cell(row, col)
                    if col == 0:
                        cell.text = str(row)
                    if col == 1:
                        cell.text = str(product_in_act[row - 1].product.product_name)
                    elif col == 2:
                        cell.text = str(product_in_act[row - 1].amount)
                        name = str(product_in_act[row - 1].product.product_name)
                        p = Product.objects.get(product_name=name)
                        p.amount = p.amount - product_in_act[row - 1].amount
                        p.save()
                    elif col == 3:
                        cell.text = str(product_in_act[row - 1].product.unit)
                    elif col == 4:
                        cell.text = str(product_in_act[row - 1].product.price_for_one)
                    elif col == 5:
                        sum = float(table.cell(row, 2).text) * float(table.cell(row, 4).text)
                        cell.text = str(sum)
                        sum_of_order = sum_of_order + sum
                elif row == r - 1 and col == 5:
                    table.cell(r - 1, 5).text = str(sum_of_order)

    doc.add_paragraph()
    doc.add_paragraph(str2)
    doc.add_paragraph(str3)

    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = 'ПРОДАВЕЦ '
    table.cell(0, 1).text = 'ПОКУПАТЕЛЬ '
    table.cell(1, 0).text = '____________/___________'
    table.cell(1, 1).text = '____________/___________'
    table.cell(2, 0).text = '      (Подпись/Ф.И.О.)'
    table.cell(2, 1).text = '      (Подпись/Ф.И.О.)'

    # root = tkinter.Tk()
    # root.withdraw()
    # root.overrideredirect(True)
    #
    # root.deiconify()
    # root.lift()
    # root.focus_force()
    # file_name = tkinter.filedialog.asksaveasfilename(initialfile=str(context['docname'] +'.docx'), filetypes=(
    #                                                             ("Doc files", ".doc;.docx"),
    #                                                             ("All files", ".*")))
    # root.destroy()
    # root.mainloop()

    # file = str(QFileDialog.getExistingDirectory(self, "Select Directory"))
    namedoc = 'C:/Users/Dell G7/Desktop/Акты/' + str(context['docname']) + '.docx'
    # if file_name:
    doc.save(namedoc)

# ЗАПОЛНЕНИЕ ФОРМЫ ДЛЯ СОЗДАНИЯ ОПИСИ
def form_opis(request):
    global CO, IDNUMOPIS
    submitbutton = request.POST.get("submit")
    object = NumOpis.objects.get(id=IDNUMOPIS)
    doc_name = ''
    podrazdel = ''
    okud = ''
    okpo = ''
    v_activ = ''
    num = ''
    date = ''
    d_start = ''
    d_end = ''
    v_oper = ''
    num_doc = str(object.numopis)
    d_make = ''
    v_product = ''
    error = ''
    forms = []
    products_in_opis = []

    for i in range(CO):
        forms.append(i)
        products_in_opis.append(i)

    form = OpisForm(request.POST or None)

    for i in range(len(forms)):
        prefix = 'form' + str(i)
        forms[i] = ProductsInOpisForm(request.POST or None, prefix=prefix)

        if form.is_valid() and forms[i].is_valid():
            podrazdel = form.cleaned_data.get('podrazdel')
            okud = form.cleaned_data.get('okud')
            okpo = form.cleaned_data.get('okpo')
            v_activ = form.cleaned_data.get('v_activ')
            num = form.cleaned_data.get('num')
            date = form.cleaned_data.get('date')
            d_start = form.cleaned_data.get('d_start')
            d_end = form.cleaned_data.get('d_end')
            v_oper = form.cleaned_data.get('v_oper')
            v_product = form.cleaned_data.get('v_product')
            d_make = str(datetime.today().strftime('"%d"   %m    %Yг.'))
            doc_name = "Инвентаризационная опись №" + str(num_doc)

            p = forms[i].cleaned_data.get('product')
            a = forms[i].cleaned_data.get('amount')

            products_in_opis[i] = ProductsInOPis(numopis=NumOpis.objects.get(id=IDNUMOPIS), product=p, amount=a)
            products_in_opis[i].save()
        else:
            error = 'Форма была неверно заполнена.'

    context = {
        'form': form, 'forms': forms, 'C': C,
        'doc_name': doc_name, 'podrazdel': podrazdel, 'okud': okud,
        'okpo': okpo, 'v_activ': v_activ, 'num': num,
        'date': date, 'd_start': d_start, 'd_end': d_end,
        'v_oper': v_oper, 'num_doc': num_doc, 'd_make': d_make,
        'v_product': v_product,
        'products_in_opis': products_in_opis,
        'error': error, 'submitbutton': submitbutton,
        'title': 'Инвентар опись'
    }

    print(context)

    if context['podrazdel'] != '':
        make_opis(context, object)

    return render(request, 'main/opis_add_form.html', context)

# СОЗДАНИЕ ДОКУМЕНТА "ИНВЕНТАРИЗАЦИОННАЯ ОПИСЬ"
def make_opis(context, numopis):
    doc = DocxTemplate('shablon2.docx')
    productf = context['products_in_opis'].copy()
    sumf = [i for i in range(len(productf))]
    sumis = [i for i in range(len(productf))]
    productsis = []
    kf = 0
    sf = 0
    kis = 0
    sis = 0

    for i in range(len(productf)):
        name = str(productf[i].product.product_name)
        p = Product.objects.get(product_name=name)
        productsis.append(p)
        print(productsis)

    l = len(productf)

    for i in range(len(productf)):
        sumf[i] = str(float(productf[i].product.price_for_one) * float(productf[i].amount))
        sumis[i] = str(float(productsis[i].price_for_one) * float(productsis[i].amount))
        kf = float(kf) + float(productf[i].amount)
        sf = float(sf) + float(sumf[i])
        kis = float(kis) + float(productsis[i].amount)
        sis = float(sis) + float(sumis[i])

    data = {
        'podrazdel': context['podrazdel'], 'okud': context['okud'],
        'okpo': context['okpo'], 'v_activ': context['v_activ'], 'num': context['num'],
        'date': context['date'], 'dstart': context['d_start'], 'dend': context['d_end'],
        'voper': context['v_oper'], 'numdoc': context['num_doc'], 'dmake': context['d_make'],
        'vproduct': context['v_product'], 'sumf': sumf, 'sumis': sumis, 'p': productf, 'ps': productsis,
        'kf': str(kf), 'sf': str(sf), 'kis': str(kis), 'sis': str(sis), 'l': l
    }
    doc.render(data)

    pathdochome = 'C:/Users/Dell G7/Desktop/Опись/' + str(context['doc_name']) + '.docx'
    path = str(context['doc_name']) + '.docx'
    doc.save(path)
    doc.save(pathdochome)

    with open(path, 'rb') as f:
            myfile = File(f)
            newdoc = Opis(docfile=myfile, date=str(datetime.today().strftime('%d.%m.%Yг.')),numopis = numopis)
            newdoc.save()

# ПРОСМОТР ОПИСЕЙ
def open_opis(request):
    search_query = request.GET.get('search', '')

    if search_query:
        documents = Opis.objects.filter(date__icontains=search_query)
    else:
        documents = Opis.objects.all()

    data = {'documents': documents}
    return render(request, 'main/open_opis.html', data)

# ПРОСМОТР АКТОВ
def open_act(request):
    search_query = request.GET.get('search', '').upper()
    message = 'Загрузите акт приема-передач.'

    if request.method == 'POST':
        form = ActOpenForm(request.POST, request.FILES)
        if form.is_valid():
            client = form.cleaned_data.get('client')
            date = form.cleaned_data.get('date')
            numact = form.cleaned_data.get('numact')

            newdoc = ActDoc(docfile=request.FILES['docfile'], client=client, date=date,
                            numact=NumAct.objects.get(numact=numact))
            newdoc.save()
            return redirect('open_act')
        else:
            message = 'The form is not valid. Fix the following error:'
    else:
        form = ActOpenForm()

    if search_query:
        documents = ActDoc.objects.filter(date__icontains=search_query)
    else:
        documents = ActDoc.objects.all()

    data = {'documents': documents, 'form': form, 'message': message,}
    return render(request, 'main/open_act.html', data)
